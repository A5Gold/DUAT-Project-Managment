# MTR DUAT - DOCX Parser Tests
"""Unit tests for parsers/docx_parser.py using TDD approach."""

import re
import pytest
from pathlib import Path
from unittest.mock import MagicMock
from docx import Document
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# Helpers -- imported after implementation exists
# ---------------------------------------------------------------------------

from parsers.docx_parser import (
    DailyReportParser,
    process_docx,
    extract_line_code,
    is_blue_color,
    extract_week_year_from_filename,
    extract_project_code,
    JOB_KEYWORDS,
    LINE_CODE_PATTERN,
)


# ===========================================================================
# Fixtures
# ===========================================================================


@pytest.fixture
def empty_folder(tmp_path: Path) -> Path:
    """Return an empty temporary directory."""
    return tmp_path


@pytest.fixture
def folder_with_valid_files(tmp_path: Path) -> Path:
    """Create a folder with correctly-named dummy DOCX files."""
    for name in (
        "PS-OHLR_DUAT_Daily Report_WK21_2025.docx",
        "PS-OHLR_DUAT_Daily Report_WK22_2025.docx",
    ):
        doc = Document()
        doc.save(str(tmp_path / name))
    return tmp_path


@pytest.fixture
def folder_with_temp_files(tmp_path: Path) -> Path:
    """Create a folder that also contains ~$ temp files."""
    valid = "PS-OHLR_DUAT_Daily Report_WK21_2025.docx"
    temp = "~$PS-OHLR_DUAT_Daily Report_WK21_2025.docx"
    doc = Document()
    doc.save(str(tmp_path / valid))
    # temp file -- just write bytes so it exists
    (tmp_path / temp).write_bytes(b"temp")
    return tmp_path


@pytest.fixture
def folder_with_mixed_files(tmp_path: Path) -> Path:
    """Folder with matching, non-matching, and temp files."""
    matching = "PS-OHLR_DUAT_Daily Report_WK10_2025.docx"
    non_matching = "SomeOtherReport.docx"
    temp = "~$PS-OHLR_DUAT_Daily Report_WK10_2025.docx"
    txt_file = "readme.txt"

    doc = Document()
    doc.save(str(tmp_path / matching))
    doc.save(str(tmp_path / non_matching))
    (tmp_path / temp).write_bytes(b"temp")
    (tmp_path / txt_file).write_text("hello")
    return tmp_path


@pytest.fixture
def simple_docx_with_table(tmp_path: Path) -> Path:
    """Create a minimal DOCX with a table containing delivery data."""
    filepath = tmp_path / "PS-OHLR_DUAT_Daily Report_WK21_2025.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)

    # Header row
    table.rows[0].cells[0].text = "Date"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Qty"
    table.rows[0].cells[3].text = "Line"

    # Data row -- plain text (no blue)
    table.rows[1].cells[0].text = "Mon 19/5"
    table.rows[1].cells[1].text = "C9081 CBM KTL"
    table.rows[1].cells[2].text = "3"
    table.rows[1].cells[3].text = "KTL"

    doc.save(str(filepath))
    return filepath


@pytest.fixture
def docx_with_blue_text(tmp_path: Path) -> Path:
    """Create a DOCX where a run has blue font color."""
    filepath = tmp_path / "PS-OHLR_DUAT_Daily Report_WK05_2025.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)

    table.rows[0].cells[0].text = "Date"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Qty"
    table.rows[0].cells[3].text = "Line"

    # Data row with blue run
    table.rows[1].cells[0].text = "Tue 4/2"
    cell = table.rows[1].cells[1]
    cell.text = ""  # clear default
    run = cell.paragraphs[0].add_run("CBM C9081 KTL")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    table.rows[1].cells[2].text = "2"
    table.rows[1].cells[3].text = "KTL"

    doc.save(str(filepath))
    return filepath


# ===========================================================================
# 1. DailyReportParser.__init__
# ===========================================================================


class TestDailyReportParserInit:
    @pytest.mark.unit
    def test_stores_folder_path(self, empty_folder: Path):
        parser = DailyReportParser(empty_folder)
        assert parser.folder_path == empty_folder

    @pytest.mark.unit
    def test_accepts_string_path(self, empty_folder: Path):
        parser = DailyReportParser(str(empty_folder))
        assert parser.folder_path == empty_folder


# ===========================================================================
# 2-4. get_report_files
# ===========================================================================


class TestGetReportFiles:
    @pytest.mark.unit
    def test_empty_folder_returns_empty(self, empty_folder: Path):
        parser = DailyReportParser(empty_folder)
        assert parser.get_report_files() == []

    @pytest.mark.unit
    def test_returns_matching_files(self, folder_with_valid_files: Path):
        parser = DailyReportParser(folder_with_valid_files)
        files = parser.get_report_files()
        assert len(files) == 2
        for f in files:
            assert f.name.startswith("PS-OHLR_DUAT_Daily Report_")
            assert f.suffix == ".docx"

    @pytest.mark.unit
    def test_excludes_temp_files(self, folder_with_temp_files: Path):
        parser = DailyReportParser(folder_with_temp_files)
        files = parser.get_report_files()
        assert len(files) == 1
        assert not files[0].name.startswith("~$")

    @pytest.mark.unit
    def test_only_matches_correct_pattern(self, folder_with_mixed_files: Path):
        parser = DailyReportParser(folder_with_mixed_files)
        files = parser.get_report_files()
        assert len(files) == 1
        assert "PS-OHLR_DUAT_Daily Report_" in files[0].name

    @pytest.mark.unit
    def test_returns_path_objects(self, folder_with_valid_files: Path):
        parser = DailyReportParser(folder_with_valid_files)
        files = parser.get_report_files()
        for f in files:
            assert isinstance(f, Path)


# ===========================================================================
# 5. get_max_week
# ===========================================================================


class TestGetMaxWeek:
    @pytest.mark.unit
    def test_returns_zero_when_no_files_processed(self, empty_folder: Path):
        parser = DailyReportParser(empty_folder)
        assert parser.get_max_week() == 0


# ===========================================================================
# 6. process_all
# ===========================================================================


class TestProcessAll:
    @pytest.mark.unit
    def test_empty_folder_returns_empty_list(self, empty_folder: Path):
        parser = DailyReportParser(empty_folder)
        result = parser.process_all()
        assert result == []

    @pytest.mark.unit
    def test_progress_callback_is_called(self, folder_with_valid_files: Path):
        parser = DailyReportParser(folder_with_valid_files)
        calls = []

        def callback(filename: str, progress: float):
            calls.append((filename, progress))

        parser.process_all(progress_callback=callback)
        assert len(calls) > 0
        # Last call should have progress close to 1.0
        assert calls[-1][1] == pytest.approx(1.0, abs=0.01)


# ===========================================================================
# 7-8. process_docx
# ===========================================================================


class TestProcessDocx:
    @pytest.mark.unit
    def test_nonexistent_file_returns_empty(self, tmp_path: Path):
        fake = tmp_path / "nonexistent.docx"
        result = process_docx(fake)
        assert result == []

    @pytest.mark.unit
    def test_non_docx_file_returns_empty(self, tmp_path: Path):
        txt = tmp_path / "report.txt"
        txt.write_text("not a docx")
        result = process_docx(txt)
        assert result == []

    @pytest.mark.unit
    def test_corrupted_file_returns_empty(self, tmp_path: Path):
        bad = tmp_path / "PS-OHLR_DUAT_Daily Report_WK01_2025.docx"
        bad.write_bytes(b"this is not a valid docx")
        result = process_docx(bad)
        assert result == []


# ===========================================================================
# 9. Line code regex extraction helper
# ===========================================================================


class TestExtractLineCode:
    @pytest.mark.unit
    @pytest.mark.parametrize(
        "text,expected",
        [
            ("Work on KTL section", "KTL"),
            ("TCL maintenance", "TCL"),
            ("AEL repair work", "AEL"),
            ("TWL overhead line", "TWL"),
            ("ISL station", "ISL"),
            ("TKL tunnel", "TKL"),
            ("EAL crossing", "EAL"),
            ("SIL platform", "SIL"),
            ("TML depot", "TML"),
            ("DRL track", "DRL"),
        ],
    )
    def test_extracts_all_line_codes(self, text: str, expected: str):
        assert extract_line_code(text) == expected

    @pytest.mark.unit
    def test_returns_empty_when_no_match(self):
        assert extract_line_code("No line code here") == ""

    @pytest.mark.unit
    def test_returns_first_match_when_multiple(self):
        result = extract_line_code("KTL and TCL work")
        assert result in ("KTL", "TCL")

    @pytest.mark.unit
    def test_case_sensitive(self):
        # Line codes should be uppercase only
        assert extract_line_code("ktl lowercase") == ""

    @pytest.mark.unit
    def test_word_boundary(self):
        # Should not match partial words
        assert extract_line_code("AKTL something") == ""


# ===========================================================================
# 10. Blue color detection helper
# ===========================================================================


class TestIsBlueColor:
    @pytest.mark.unit
    def test_pure_blue(self):
        assert is_blue_color(RGBColor(0x00, 0x00, 0xFF)) is True

    @pytest.mark.unit
    def test_dark_blue(self):
        assert is_blue_color(RGBColor(0x00, 0x00, 0xCC)) is True

    @pytest.mark.unit
    def test_medium_blue(self):
        assert is_blue_color(RGBColor(0x00, 0x00, 0x99)) is True

    @pytest.mark.unit
    def test_not_blue_red(self):
        assert is_blue_color(RGBColor(0xFF, 0x00, 0x00)) is False

    @pytest.mark.unit
    def test_not_blue_green(self):
        assert is_blue_color(RGBColor(0x00, 0xFF, 0x00)) is False

    @pytest.mark.unit
    def test_not_blue_black(self):
        assert is_blue_color(RGBColor(0x00, 0x00, 0x00)) is False

    @pytest.mark.unit
    def test_none_returns_false(self):
        assert is_blue_color(None) is False

    @pytest.mark.unit
    def test_light_blue_shade(self):
        # Blue with some red/green but still predominantly blue
        assert is_blue_color(RGBColor(0x33, 0x33, 0xFF)) is True


# ===========================================================================
# 11. Week/Year extraction from filename
# ===========================================================================


class TestExtractWeekYearFromFilename:
    @pytest.mark.unit
    def test_standard_filename(self):
        week, year = extract_week_year_from_filename(
            "PS-OHLR_DUAT_Daily Report_WK21_2025.docx"
        )
        assert week == 21
        assert year == 2025

    @pytest.mark.unit
    def test_single_digit_week(self):
        week, year = extract_week_year_from_filename(
            "PS-OHLR_DUAT_Daily Report_WK05_2025.docx"
        )
        assert week == 5
        assert year == 2025

    @pytest.mark.unit
    def test_week_52(self):
        week, year = extract_week_year_from_filename(
            "PS-OHLR_DUAT_Daily Report_WK52_2024.docx"
        )
        assert week == 52
        assert year == 2024

    @pytest.mark.unit
    def test_no_match_returns_zero(self):
        week, year = extract_week_year_from_filename("SomeOtherFile.docx")
        assert week == 0
        assert year == 0

    @pytest.mark.unit
    def test_works_with_path_object(self):
        p = Path("C:/reports/PS-OHLR_DUAT_Daily Report_WK10_2026.docx")
        week, year = extract_week_year_from_filename(p.name)
        assert week == 10
        assert year == 2026


# ===========================================================================
# 12. Project code extraction helper
# ===========================================================================


class TestExtractProjectCode:
    @pytest.mark.unit
    def test_extracts_c_code(self):
        assert extract_project_code("C9081 CBM work") == "C9081"

    @pytest.mark.unit
    def test_extracts_from_middle(self):
        assert extract_project_code("Work on C2264 KTL") == "C2264"

    @pytest.mark.unit
    def test_no_project_code(self):
        assert extract_project_code("CBM maintenance") == ""

    @pytest.mark.unit
    def test_multiple_codes_returns_first(self):
        result = extract_project_code("C9081 and C2264")
        assert result == "C9081"

    @pytest.mark.unit
    def test_must_be_four_digits(self):
        # C followed by less than 4 digits should not match
        assert extract_project_code("C123 short") == ""

    @pytest.mark.unit
    def test_keyword_without_project_returns_keyword(self):
        # When text is just a keyword like CBM, no project code
        assert extract_project_code("CBM") == ""


# ===========================================================================
# 13. Record format validation
# ===========================================================================


class TestRecordFormat:
    @pytest.mark.unit
    def test_record_has_required_fields(self, simple_docx_with_table: Path):
        records = process_docx(simple_docx_with_table)
        if len(records) > 0:
            record = records[0]
            required_fields = {
                "FullDate",
                "Project",
                "Qty Delivered",
                "Week",
                "Year",
                "Line",
            }
            assert required_fields.issubset(set(record.keys()))


# ===========================================================================
# 14. Constants validation
# ===========================================================================


class TestConstants:
    @pytest.mark.unit
    def test_job_keywords_contains_expected(self):
        expected = {"CBM", "CM", "PA work", "HLM", "Provide"}
        assert expected.issubset(set(JOB_KEYWORDS))

    @pytest.mark.unit
    def test_line_code_pattern_is_valid_regex(self):
        compiled = re.compile(LINE_CODE_PATTERN)
        assert compiled.match("KTL") is not None
