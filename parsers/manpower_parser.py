# MTR PS-OHLR DUAT - Manpower Parser
"""
Parse manpower/shift data from the SECOND table of daily report DOCX files.

Extracts shift records including jobs, team counts, EPIC roles,
attendance, and leave information.
"""

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ── Constants ─────────────────────────────────────────────────────────────────

REPORT_GLOB = "PS-OHLR_DUAT_Daily Report_*.docx"

TEAM_COUNT_RE = re.compile(r"S([2-5])x(\d+)")
PROJECT_CODE_RE = re.compile(r"C\d{4}")

JOB_TYPE_KEYWORDS: List[Tuple[str, str]] = [
    ("SPA work", "SPA work"),
    ("PA work", "PA work"),
    ("CBM", "CBM"),
    ("CM", "CM"),
    ("HLM", "HLM"),
    ("C&R", "C&R"),
]

LEAVE_CATEGORIES = ("AL", "SH", "SL", "RD", "Training")

ROLE_PATTERNS: List[Tuple[str, str]] = [
    (r"CP\s*\(\s*P\s*\)", "CP_P"),
    (r"CP\s*\(\s*T\s*\)", "CP_T"),
    (r"AP\s*\(\s*E\s*\)", "AP_E"),
    (r"SPC", "SPC"),
    (r"HSM", "HSM"),
    (r"NP", "NP"),
]

EMPTY_LEAVE: Dict[str, List[str]] = {
    "AL": [], "SH": [], "SL": [], "RD": [], "Training": [],
}

EMPTY_ROLES: Dict = {
    "EPIC": "",
    "CP_P": [],
    "CP_T": [],
    "AP_E": [],
    "SPC": [],
    "HSM": [],
    "NP": [],
}


# ── Helper Functions ──────────────────────────────────────────────────────────


def _parse_team_counts(text: str) -> Dict[str, int]:
    """
    Parse team count string like 'S2x3, S3x2' into {S2: 3, S3: 2}.

    Uses regex S[2-5]x\\d+ to find all team assignments.
    """
    counts: Dict[str, int] = {}
    for match in TEAM_COUNT_RE.finditer(text):
        team_id = f"S{match.group(1)}"
        count = int(match.group(2))
        counts[team_id] = count
    return counts


def _extract_project_code(text: str) -> Optional[str]:
    """
    Extract first project code (C####) from description text.

    Returns None if no project code found.
    """
    match = PROJECT_CODE_RE.search(text)
    return match.group(0) if match else None


def _classify_job_type(text: str) -> str:
    """
    Classify job type from description text using keyword matching.

    Order matters: SPA work must be checked before PA work.
    Returns 'Other' if no keyword matches.
    """
    for keyword, job_type in JOB_TYPE_KEYWORDS:
        if keyword in text:
            return job_type
    return "Other"


def _categorize_leave(text: str) -> Dict[str, List[str]]:
    """
    Parse leave text into categories: AL, SH, SL, RD, Training.

    Expected format per line: 'CATEGORY: name1, name2'
    """
    result: Dict[str, List[str]] = {cat: [] for cat in LEAVE_CATEGORIES}
    if not text.strip():
        return result

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        for category in LEAVE_CATEGORIES:
            pattern = f"{category}:"
            if pattern in line:
                names_part = line.split(pattern, 1)[1].strip()
                names = [
                    n.strip() for n in names_part.split(",") if n.strip()
                ]
                result[category].extend(names)
                break

    return result


def _parse_epic_roles(text: str) -> Dict:
    """
    Extract EPIC safety roles from role assignment text.

    Looks for patterns like 'CP(P): John', 'AP(E): Tom, Mary'.
    Returns dict with EPIC string and lists per role key.
    """
    roles: Dict = {
        "EPIC": "",
        "CP_P": [],
        "CP_T": [],
        "AP_E": [],
        "SPC": [],
        "HSM": [],
        "NP": [],
    }

    if not text.strip():
        return roles

    roles["EPIC"] = text.strip()

    for pattern, role_key in ROLE_PATTERNS:
        role_re = re.compile(
            pattern + r"\s*:\s*([^\n]+?)(?=\s*(?:"
            + "|".join(p for p, _ in ROLE_PATTERNS)
            + r")\s*:|$)",
            re.IGNORECASE,
        )
        for match in role_re.finditer(text):
            names_str = match.group(1).strip().rstrip(",")
            names = [n.strip() for n in names_str.split(",") if n.strip()]
            roles[role_key].extend(names)

    return roles


def _extract_names_from_text(text: str) -> List[str]:
    """Extract comma-separated names from text, filtering empty strings."""
    if not text.strip():
        return []
    return [n.strip() for n in text.split(",") if n.strip()]


def _parse_shift_from_filename(filename: str) -> Tuple[str, str]:
    """
    Extract week and year hints from filename.

    Filename pattern: PS-OHLR_DUAT_Daily Report_*.docx
    Returns (week, year) as strings, or empty strings if not found.
    """
    week = ""
    year = ""

    wk_match = re.search(r"WK(\d+)", filename, re.IGNORECASE)
    if wk_match:
        week = wk_match.group(1)

    year_match = re.search(r"20(\d{2})", filename)
    if year_match:
        year = f"20{year_match.group(1)}"

    return week, year


# ── Table Parsing ─────────────────────────────────────────────────────────────


def _get_cell_text(cell) -> str:
    """Get full text content of a table cell, joining paragraphs."""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _parse_second_table(table, week: str, year: str) -> List[Dict]:
    """
    Parse the second table of a daily report DOCX for shift records.

    The table contains shift information split into sections:
    HLM, C&R Works and Projects, Attendance, Other Notable Items.
    """
    records: List[Dict] = []
    rows = table.rows

    if len(rows) < 2:
        return records

    current_date = ""
    current_day = ""
    current_shift = ""

    for row in rows[1:]:
        cells = row.cells
        if len(cells) < 2:
            continue

        cell_texts = [_get_cell_text(c) for c in cells]
        first_cell = cell_texts[0] if cell_texts else ""

        # Detect date pattern
        date_match = re.search(r"(\d{1,2}/\d{1,2}(?:/\d{2,4})?)", first_cell)
        if date_match:
            current_date = date_match.group(1)
            day_match = re.search(
                r"(Mon|Tue|Wed|Thu|Fri|Sat|Sun)\w*",
                first_cell,
                re.IGNORECASE,
            )
            if day_match:
                current_day = day_match.group(0)

        # Detect shift indicator
        for cell_text in cell_texts:
            cell_lower = cell_text.lower()
            if "day" in cell_lower and "shift" in cell_lower:
                current_shift = "Day"
            elif "night" in cell_lower and "shift" in cell_lower:
                current_shift = "Night"
            elif cell_text.strip() == "Day":
                current_shift = "Day"
            elif cell_text.strip() == "Night":
                current_shift = "Night"

        full_row_text = " ".join(cell_texts)
        lower_text = full_row_text.lower()

        # Parse job entries
        job_type = _classify_job_type(full_row_text)
        has_project = _extract_project_code(full_row_text) is not None

        if job_type != "Other" or has_project:
            project_code = _extract_project_code(full_row_text)
            team_counts = _parse_team_counts(full_row_text)

            qty = 0.0
            for ct in cell_texts:
                try:
                    val = float(ct.strip())
                    qty = val
                    break
                except (ValueError, AttributeError):
                    continue

            roles = _parse_epic_roles(full_row_text)
            done_by_raw = ""
            for ct in cell_texts:
                if TEAM_COUNT_RE.search(ct):
                    done_by_raw = ct.strip()
                    break

            worker_names = _extract_names_from_text(done_by_raw)
            total_workers = sum(team_counts.values()) + len(worker_names)

            resolved_type = job_type if job_type != "Other" else (
                "C&R" if project_code else "Other"
            )

            job: Dict = {
                "type": resolved_type,
                "project_code": project_code,
                "description": full_row_text[:200],
                "qty": qty,
                "done_by_raw": done_by_raw,
                "team_counts": team_counts,
                "worker_names": worker_names,
                "total_workers": total_workers,
                "roles": roles,
            }

            needs_new_record = (
                not records
                or records[-1]["date"] != current_date
                or records[-1]["shift"] != (current_shift or "Day")
            )
            if needs_new_record:
                record: Dict = {
                    "date": current_date,
                    "day_of_week": current_day,
                    "shift": current_shift or "Day",
                    "week": week,
                    "year": year,
                    "jobs": [job],
                    "on_duty_names": [],
                    "on_duty_team_counts": {},
                    "apprentices": [],
                    "term_labour_count": 0,
                    "leave": {cat: [] for cat in LEAVE_CATEGORIES},
                }
                records.append(record)
            else:
                records[-1]["jobs"].append(job)

        # Parse attendance
        if "on duty" in lower_text or "attendance" in lower_text:
            after_colon = (
                full_row_text.split(":", 1)[-1] if ":" in full_row_text else ""
            )
            names = _extract_names_from_text(after_colon)
            team_counts = _parse_team_counts(full_row_text)
            if records:
                records[-1]["on_duty_names"] = names
                records[-1]["on_duty_team_counts"] = team_counts

        # Parse apprentices
        if "apprentice" in lower_text:
            after_colon = (
                full_row_text.split(":", 1)[-1] if ":" in full_row_text else ""
            )
            names = _extract_names_from_text(after_colon)
            if records:
                records[-1]["apprentices"] = names

        # Parse term labour
        if "term labour" in lower_text or "term labor" in lower_text:
            count_src = (
                full_row_text.split(":")[-1] if ":" in full_row_text
                else full_row_text
            )
            count_match = re.search(r"(\d+)", count_src)
            if count_match and records:
                records[-1]["term_labour_count"] = int(count_match.group(1))

        # Parse leave
        for category in LEAVE_CATEGORIES:
            if f"{category}:" in full_row_text:
                leave_data = _categorize_leave(full_row_text)
                if records:
                    for cat_key, cat_names in leave_data.items():
                        if cat_names:
                            records[-1]["leave"][cat_key].extend(cat_names)
                break

    return records


# ── ManpowerParser Class ──────────────────────────────────────────────────────


class ManpowerParser:
    """
    Parse manpower data from daily report DOCX files.

    Reads the SECOND table of each DOCX file matching the report pattern,
    extracting shift records with jobs, attendance, and leave data.
    """

    def __init__(self, folder_path: Path) -> None:
        self.folder_path = Path(folder_path)

    def get_report_files(self) -> List[Path]:
        """
        Find all matching daily report DOCX files in the folder.

        Excludes temporary files starting with '~$'.
        Returns sorted list of Path objects.
        """
        if not self.folder_path.exists():
            return []

        files = [
            f for f in self.folder_path.glob(REPORT_GLOB)
            if not f.name.startswith("~$")
        ]
        return sorted(files, key=lambda p: p.name)

    def process_all(self) -> List[Dict]:
        """
        Parse all report files and return list of ShiftRecord dicts.

        Corrupted files are logged and skipped gracefully.
        """
        from docx import Document

        files = self.get_report_files()
        if not files:
            return []

        all_records: List[Dict] = []

        for filepath in files:
            try:
                doc = Document(str(filepath))
                tables = doc.tables

                if len(tables) < 2:
                    logger.warning(
                        "File %s has fewer than 2 tables, skipping",
                        filepath.name,
                    )
                    continue

                week, year = _parse_shift_from_filename(filepath.name)
                records = _parse_second_table(tables[1], week, year)
                all_records.extend(records)

                logger.info(
                    "Parsed %d records from %s",
                    len(records),
                    filepath.name,
                )

            except Exception as exc:
                logger.error(
                    "Failed to parse %s: %s",
                    filepath.name,
                    exc,
                )
                continue

        return all_records
