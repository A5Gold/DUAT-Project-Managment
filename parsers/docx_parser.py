# MTR DUAT - DOCX Parser
"""
Parse daily report DOCX files and extract delivery records.

Each record contains: FullDate, Project, Qty Delivered, Week, Year, Line.
"""

import logging
import re
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import RGBColor

logger = logging.getLogger("duat.parser")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

LINE_CODE_PATTERN = r"\b(KTL|TCL|AEL|TWL|ISL|TKL|EAL|SIL|TML|DRL)\b"

JOB_KEYWORDS: List[str] = ["CBM", "CM", "PA work", "HLM", "Provide"]

_FILENAME_PATTERN = re.compile(
    r"PS-OHLR_DUAT_Daily Report_WK(\d{1,2})_(\d{4})\.docx$",
    re.IGNORECASE,
)

_PROJECT_CODE_PATTERN = re.compile(r"\bC(\d{4})\b")

_BLUE_THRESHOLD_BLUE = 0x80
_BLUE_THRESHOLD_OTHER = 0x80

# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------


def is_blue_color(rgb: Optional[RGBColor]) -> bool:
    """Return True when *rgb* represents a blue shade.

    Blue is defined as: blue channel >= 0x80 **and** both red and green
    channels are below 0x80.  ``None`` is treated as "no colour" (not blue).
    """
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return b >= _BLUE_THRESHOLD_BLUE and r < _BLUE_THRESHOLD_OTHER and g < _BLUE_THRESHOLD_OTHER


def extract_line_code(text: str) -> str:
    """Extract the first railway line code from *text*.

    Returns the matched code (e.g. ``"KTL"``) or ``""`` when none found.
    """
    match = re.search(LINE_CODE_PATTERN, text)
    return match.group(1) if match else ""


def extract_project_code(text: str) -> str:
    """Extract the first C#### project code from *text*.

    Returns the full code (e.g. ``"C9081"``) or ``""`` when none found.
    """
    match = _PROJECT_CODE_PATTERN.search(text)
    return match.group(0) if match else ""


def extract_week_year_from_filename(filename: str) -> Tuple[int, int]:
    """Parse week and year from a daily-report filename.

    Expected pattern: ``PS-OHLR_DUAT_Daily Report_WK##_YYYY.docx``

    Returns ``(week, year)`` or ``(0, 0)`` when the pattern does not match.
    """
    match = _FILENAME_PATTERN.search(filename)
    if match:
        return int(match.group(1)), int(match.group(2))
    return 0, 0


# ---------------------------------------------------------------------------
# Cell / run helpers
# ---------------------------------------------------------------------------


def _cell_has_blue_text(cell) -> bool:
    """Return True if any run inside *cell* has blue font colour."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.color and is_blue_color(run.font.color.rgb):
                return True
    return False


def _cell_text(cell) -> str:
    """Return the full plain-text content of a table cell."""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _is_night_shift_row(row_text: str) -> bool:
    """Heuristic: row belongs to a night shift section."""
    lower = row_text.lower()
    return "night" in lower


def _has_job_keyword(text: str) -> bool:
    """Return True if *text* contains any recognised job keyword."""
    for kw in JOB_KEYWORDS:
        if kw.lower() in text.lower():
            return True
    return False


def _extract_qty(text: str) -> float:
    """Try to parse a numeric quantity from *text*.

    Returns 0.0 when the text cannot be interpreted as a number.
    """
    cleaned = text.strip()
    if not cleaned:
        return 0.0
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def _identify_project(text: str) -> str:
    """Determine the Project value from cell text.

    Priority:
    1. C#### project code
    2. First matching job keyword
    3. Empty string
    """
    code = extract_project_code(text)
    if code:
        return code
    for kw in JOB_KEYWORDS:
        if kw.lower() in text.lower():
            return kw
    return ""


# ---------------------------------------------------------------------------
# Core DOCX processing
# ---------------------------------------------------------------------------


def _parse_table_records(
    table,
    week: str,
    year: str,
    is_night: bool = False,
) -> List[Dict]:
    """Extract delivery records from a single DOCX table.

    Skips the header row (index 0).  For each subsequent row, attempts to
    build a record from the cell contents.
    """
    records: List[Dict] = []
    rows = table.rows

    if len(rows) < 2:
        return records

    for row in rows[1:]:
        cells = row.cells
        if len(cells) < 2:
            continue

        # Gather full row text for context
        row_text = " ".join(_cell_text(c) for c in cells)
        if not row_text.strip():
            continue

        # Detect night shift from row context
        night = is_night or _is_night_shift_row(row_text)

        # Check for blue text in any cell of this row
        has_blue = any(_cell_has_blue_text(c) for c in cells)

        # Try to extract a date from the first cell
        full_date = _cell_text(cells[0])

        # Description cell (second cell typically)
        desc_text = _cell_text(cells[1]) if len(cells) > 1 else ""

        # Qty cell (third cell if available)
        qty_text = _cell_text(cells[2]) if len(cells) > 2 else ""
        qty = _extract_qty(qty_text)

        # Determine project
        combined_text = f"{desc_text} {row_text}"
        project = _identify_project(combined_text)

        if not project and not full_date:
            continue

        # Night shift + blue keyword -> qty = 0
        if night and has_blue and _has_job_keyword(combined_text):
            qty = 0.0

        # Extract line code from all available text
        line = extract_line_code(combined_text)

        if project or full_date:
            record: Dict = {
                "FullDate": full_date,
                "Project": project,
                "Qty Delivered": qty,
                "Week": week,
                "Year": year,
                "Line": line,
            }
            records.append(record)

    return records


def process_docx(filepath: Path) -> List[Dict]:
    """Parse a single DOCX file and return a list of delivery records.

    Each record is a dict with keys:
    ``FullDate``, ``Project``, ``Qty Delivered``, ``Week``, ``Year``, ``Line``.

    Returns an empty list when the file does not exist, is not a ``.docx``,
    or cannot be opened.
    """
    filepath = Path(filepath)

    if not filepath.exists():
        logger.warning("File does not exist: %s", filepath)
        return []

    if filepath.suffix.lower() != ".docx":
        logger.warning("Not a .docx file: %s", filepath)
        return []

    # Extract week/year from filename
    week_num, year_num = extract_week_year_from_filename(filepath.name)
    week = str(week_num) if week_num else ""
    year = str(year_num) if year_num else ""

    try:
        doc = Document(str(filepath))
    except Exception as exc:
        logger.error("Failed to open DOCX %s: %s", filepath, exc)
        return []

    records: List[Dict] = []
    for table in doc.tables:
        table_records = _parse_table_records(table, week, year)
        records.extend(table_records)

    return records


# ---------------------------------------------------------------------------
# DailyReportParser class
# ---------------------------------------------------------------------------


class DailyReportParser:
    """Scan a folder of daily-report DOCX files and extract delivery records.

    Usage::

        parser = DailyReportParser(Path("C:/Reports"))
        files = parser.get_report_files()
        records = parser.process_all(progress_callback=my_cb)
        max_week = parser.get_max_week()
    """

    def __init__(self, folder_path) -> None:
        self.folder_path: Path = Path(folder_path)
        self._records: List[Dict] = []
        self._max_week: int = 0

    # -- public API ---------------------------------------------------------

    def get_report_files(self) -> List[Path]:
        """Return sorted list of matching DOCX report files.

        Matches ``PS-OHLR_DUAT_Daily Report_*.docx`` and excludes
        temporary files whose names start with ``~$``.
        """
        if not self.folder_path.exists() or not self.folder_path.is_dir():
            return []

        matched: List[Path] = []
        for f in self.folder_path.glob("PS-OHLR_DUAT_Daily Report_*.docx"):
            if f.name.startswith("~$"):
                continue
            matched.append(f)

        return sorted(matched)

    def process_all(
        self,
        progress_callback: Optional[Callable[[str, float], None]] = None,
    ) -> List[Dict]:
        """Parse every matching file and return all records.

        *progress_callback*, when provided, is called with
        ``(filename, progress_float)`` after each file is processed.
        ``progress_float`` ranges from 0.0 to 1.0.
        """
        files = self.get_report_files()
        if not files:
            return []

        all_records: List[Dict] = []
        total = len(files)

        for idx, fpath in enumerate(files):
            file_records = process_docx(fpath)
            all_records.extend(file_records)

            # Track max week
            wk, _ = extract_week_year_from_filename(fpath.name)
            if wk > self._max_week:
                self._max_week = wk

            if progress_callback is not None:
                progress = (idx + 1) / total
                progress_callback(fpath.name, progress)

        self._records = all_records
        return all_records

    def get_max_week(self) -> int:
        """Return the highest week number seen across processed files.

        Returns ``0`` if no files have been processed yet.
        """
        return self._max_week
