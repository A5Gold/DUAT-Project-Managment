# MTR DUAT - Keyword Search Router
"""Keyword search API endpoints for scanning daily report DOCX files."""

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from pathlib import Path
import re
from docx import Document
import logging

logger = logging.getLogger(__name__)

router = APIRouter()

# Import shared state from services
from backend.services import search_state


class KeywordSearchRequest(BaseModel):
    folder_path: str
    keyword: str


@router.post("/search")
async def search_keyword(request: KeywordSearchRequest):
    """Search for a keyword across all daily report DOCX files."""
    folder = Path(request.folder_path)
    if not folder.exists():
        raise HTTPException(status_code=404, detail="Folder path does not exist")

    keyword = request.keyword.strip()
    if not keyword:
        raise HTTPException(status_code=400, detail="Keyword cannot be empty")

    docx_files = sorted([
        f for f in folder.glob("PS-OHLR_DUAT_Daily Report_*.docx")
        if not f.name.startswith("~$")
    ])

    total_files = len(docx_files)
    all_matches = []
    matched_file_count = 0

    for filepath in docx_files:
        try:
            doc = Document(str(filepath))
            file_matches = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text and keyword.lower() in text.lower():
                    file_matches.append({"location": "Paragraph", "text": text[:500]})

            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text and keyword.lower() in cell_text.lower():
                            file_matches.append({
                                "location": f"Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}",
                                "text": cell_text[:500]
                            })

            # Deduplicate
            seen = set()
            unique = []
            for m in file_matches:
                if m["text"] not in seen:
                    seen.add(m["text"])
                    unique.append(m)

            if unique:
                matched_file_count += 1
                all_matches.append({"filename": filepath.name, "matches": unique})
        except Exception as e:
            logger.warning("Failed to process %s: %s", filepath.name, e)
            continue

    search_state.update({
        "results": all_matches, "keyword": keyword,
        "total_files": total_files, "matched_files": matched_file_count
    })

    return {
        "keyword": keyword,
        "total_files": total_files,
        "matched_files": matched_file_count,
        "results": all_matches
    }
